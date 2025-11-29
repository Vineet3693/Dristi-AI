"""PDF and DOCX export functionality."""

from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.colors import HexColor
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from pathlib import Path
import re
from typing import List, Dict


class ExportHandler:
    """Handle PDF and DOCX export with proper formatting."""
    
    def __init__(self):
        """Initialize export handler."""
        self.golden = HexColor('#FFD700')
        self.orange = HexColor('#FFA500')
        self.saffron = HexColor('#FF6600')
        self.dark_gray = HexColor('#333333')
    
    def export_to_pdf(self, content: str, filename: str) -> str:
        """
        Export content to PDF with formatting.
        
        Args:
            content: Markdown content to export
            filename: Output filename
            
        Returns:
            Path to created PDF file
        """
        output_path = Path(filename)
        
        # Create PDF
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=self.golden,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=self.golden,
            spaceAfter=12
        )
        
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=12,
            textColor=self.dark_gray,
            spaceAfter=12
        )
        
        citation_style = ParagraphStyle(
            'Citation',
            parent=styles['BodyText'],
            fontSize=11,
            textColor=self.orange,
            leftIndent=20,
            spaceAfter=12
        )
        
        # Parse content
        story = []
        
        # Add header
        story.append(Paragraph("Drishti AI - Divine Wisdom", title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Parse markdown
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            
            if not line:
                story.append(Spacer(1, 0.1 * inch))
                continue
            
            # Headings
            if line.startswith('###'):
                text = line.replace('###', '').strip()
                story.append(Paragraph(text, heading_style))
            elif line.startswith('##'):
                text = line.replace('##', '').strip()
                story.append(Paragraph(text, heading_style))
            elif line.startswith('#'):
                text = line.replace('#', '').strip()
                story.append(Paragraph(text, title_style))
            
            # Citations
            elif line.startswith('**Bhagavad Gita'):
                story.append(Paragraph(line, citation_style))
            
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                text = '• ' + line[2:]
                story.append(Paragraph(text, body_style))
            
            # Regular text
            else:
                story.append(Paragraph(line, body_style))
        
        # Build PDF
        doc.build(story)
        
        return str(output_path)
    
    def export_to_docx(self, content: str, filename: str) -> str:
        """
        Export content to DOCX with formatting.
        
        Args:
            content: Markdown content to export
            filename: Output filename
            
        Returns:
            Path to created DOCX file
        """
        output_path = Path(filename)
        
        # Create document
        doc = Document()
        
        # Add header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "Drishti AI - Divine Wisdom from Bhagavad Gita"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parse content
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            
            if not line:
                doc.add_paragraph()
                continue
            
            # Headings
            if line.startswith('###'):
                text = line.replace('###', '').strip()
                para = doc.add_heading(text, level=3)
                para.runs[0].font.color.rgb = RGBColor(255, 215, 0)
            
            elif line.startswith('##'):
                text = line.replace('##', '').strip()
                para = doc.add_heading(text, level=2)
                para.runs[0].font.color.rgb = RGBColor(255, 215, 0)
            
            elif line.startswith('#'):
                text = line.replace('#', '').strip()
                para = doc.add_heading(text, level=1)
                para.runs[0].font.color.rgb = RGBColor(255, 215, 0)
            
            # Citations
            elif line.startswith('**Bhagavad Gita'):
                para = doc.add_paragraph(line)
                para.runs[0].font.color.rgb = RGBColor(255, 165, 0)
                para.runs[0].font.size = Pt(11)
            
            # Bullet points
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:]
                doc.add_paragraph(text, style='List Bullet')
            
            # Regular text
            else:
                para = doc.add_paragraph(line)
                para.runs[0].font.size = Pt(12)
        
        # Save document
        doc.save(str(output_path))
        
        return str(output_path)
